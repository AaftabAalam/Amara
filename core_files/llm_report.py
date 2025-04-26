from langchain_groq import ChatGroq
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.chains import ConversationalRetrievalChain
from fastapi import File, UploadFile, HTTPException
from docx import Document
import uuid
import csv
import pdfplumber
import os
from typing import Union, List
from fastapi import File, UploadFile
from typing import List
import os
from pydantic import BaseModel
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from dotenv import load_dotenv
load_dotenv()

PERSIST_DIRECTORY = "db1"
os.makedirs(PERSIST_DIRECTORY, exist_ok=True)

embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

def get_vector_store(user_id: str):
    return Chroma(
        persist_directory=f"{PERSIST_DIRECTORY}/{user_id}",
        embedding_function=embeddings
    )

class QueryRequest(BaseModel):
    user_id: str
    question: str

from langchain.schema import Document

async def upload_document(user_id: str, files: List[UploadFile]):
    try:
        user_data = f"{PERSIST_DIRECTORY}/{user_id}"
        os.makedirs(user_data, exist_ok=True)

        try:
            vector_store = get_vector_store(user_id)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error initializing vector store: {str(e)}")

        for file in files:
            try:
                file_extension = file.filename.split('.')[-1].lower()

                if file_extension not in ['pdf', 'docx', 'txt', 'csv']:  
                    return {"error": f"Unsupported file type: {file_extension}"}
            
                file_path = f"{user_data}/{str(uuid.uuid4())}_{file.filename}"

                with open(file_path, "wb") as f:
                    f.write(await file.read())

                try:
                    document_text = extract_text_from_file(file_path)
                except Exception as e:
                    raise HTTPException(status_code=400, detail=f"Error extracting text from file: {str(e)}")

                document = Document(
                    page_content=document_text,
                    metadata={"source": file.filename, "file_path": file_path}
                )
                try:
                    vector_store.add_documents([document])
                except Exception as e:
                    raise HTTPException(status_code=500, detail=f"Error storing document to vector store: {str(e)}")
            except Exception as e:
                return {"error": f"Failed to process files '{file.filename}': {str(e)}"}
        return {"message": "Document uploaded successfully and embeddings stored."}
    except Exception as e:
        return {"error": f"An unexpected error occured: {str(e)}"}

async def ask_query(query: QueryRequest):
    try:
        try:
            vector_store = get_vector_store(query.user_id)
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error retrieving vector store for user_id {query.user_id}: {str(e)}"
            )

        try:
            chat_model = ChatGroq(
                groq_api_key=os.getenv("GROQ_API_KEY"),
                model_name="gemma2-9b-it"
            )
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error initializing chat model: {str(e)}"
            )

        try:
            qa_chain = ConversationalRetrievalChain.from_llm(
                llm=chat_model,
                retriever=vector_store.as_retriever(),
                return_source_documents=True
            )
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error initializing QA chain: {str(e)}"
            )

        try:
            response = qa_chain({"question": query.question, "chat_history": []})
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error during query processing: {str(e)}"
            )

        if not response or response.get("answer", "").strip() == "":
            return {"message": "No relevant information found."}

        return {
            'answer': response['answer'],
            'sources': [str(doc.metadata) for doc in response.get('source_documents', [])]
        }
    except HTTPException as http_exc:
        raise http_exc
    except Exception as e:
        return {"error": f"An unexpected error occurred: {str(e)}"}

def extract_text_from_file(file_path: str) -> Union[str, None]:
    try:
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == ".pdf":
            try:
                with pdfplumber.open(file_path) as pdf:
                    text = ""
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except Exception as e:
                raise ValueError(f"Error extracting text from PDF: {str(e)}")
        
        elif file_extension == ".docx":
            try:
                doc = Document(file_path)
                text = ""
                for para in doc.paragraphs:
                    text += para.text + "\n"
                return text
            except Exception as e:
                raise ValueError(f"Error extracting text from DOCX: {str(e)}")
        
        elif file_extension == ".txt":
            try:
                with open(file_path, "r", encoding="utf-8") as file:
                    text = file.read()
                return text
            except Exception as e:
                raise ValueError(f"Error reading text file: {str(e)}")
        
        elif file_extension == ".csv":
            try:
                text = ""
                with open(file_path, newline="", encoding="utf-8") as file:
                    reader = csv.reader(file)
                    for row in reader:
                        text += " ".join(row) + "\n"
                return text
            except Exception as e:
                raise ValueError(f"Error reading CSV file: {str(e)}")
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

    except ValueError as ve:
        raise ve
    except Exception as e:
        raise ValueError(f"An unexpected error occurred while extracting text: {str(e)}")
